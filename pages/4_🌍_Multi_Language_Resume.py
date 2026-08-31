import io
import streamlit as st
from google import genai

st.set_page_config(page_title="Multi-Language Resume", page_icon="🌍", layout="wide")

LANGUAGES = {
    "English": "English",
    "French": "French",
    "Spanish": "Spanish",
    "German": "German",
    "Italian": "Italian",
    "Portuguese": "Portuguese",
    "Arabic": "Arabic",
    "Amharic": "Amharic",
    "Hindi": "Hindi",
    "Chinese (Simplified)": "Simplified Chinese",
    "Japanese": "Japanese",
    "Turkish": "Turkish",
}

MODEL = "gemini-2.5-flash"


def get_client():
    key = st.secrets.get("GOOGLE_API_KEY")
    if not key:
        st.error("Missing `GOOGLE_API_KEY` in `.streamlit/secrets.toml`. Add it and restart the app.")
        st.stop()
    return genai.Client(api_key=key)


def translate_resume(client, resume_text: str, target_language: str) -> str:
    prompt = f"""You are a professional resume translator.
Translate the resume below into {target_language}.

Rules:
- Keep the original structure, section headings, bullet points, dates and numbers.
- Keep proper nouns (company names, degrees, certifications) unchanged unless there is a standard local translation.
- Keep technical terms (e.g. Python, React, ATS) unchanged.
- Output ONLY the translated resume as plain text.

RESUME:
{resume_text}"""
    resp = client.models.generate_content(model=MODEL, contents=prompt)
    return resp.text.strip()


def rewrite_in_language(client, resume_text: str, target_language: str, focus: str = "") -> str:
    prompt = f"""Rewrite and improve the following resume so it reads natively in {target_language}.
Keep all factual information identical (dates, companies, roles, skills) but improve wording, action verbs and impact.
{focus}
Output ONLY the improved resume in {target_language}, as plain text.

RESUME:
{resume_text}"""
    resp = client.models.generate_content(model=MODEL, contents=prompt)
    return resp.text.strip()


def to_docx(text: str) -> bytes:
    from docx import Document
    doc = Document()
    for line in text.splitlines():
        if not line.strip():
            continue
        if len(line.strip()) < 80:
            doc.add_heading(line.strip(), level=2)
        else:
            doc.add_paragraph(line.strip())
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def to_pdf(text: str) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.pdfgen import canvas
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    width, height = A4
    y = height - 2 * cm
    for line in text.splitlines():
        if y < 2 * cm:
            c.showPage()
            y = height - 2 * cm
        c.drawString(2 * cm, y, line[:110])
        y -= 0.6 * cm
    c.save()
    buf.seek(0)
    return buf.getvalue()


st.title("🌍 Multi-Language Resume")
st.caption("Translate your resume, or regenerate it natively in another language — then download as Word or PDF.")

col_left, col_right = st.columns([1, 1], gap="large")

with col_left:
    st.subheader("1. Paste your resume text")
    resume_text = st.text_area("Resume text", height=300, placeholder="Paste your parsed resume text here…")

with col_right:
    st.subheader("2. Choose language & action")
    target_language = st.selectbox("Target language", list(LANGUAGES.keys()), index=0)
    action = st.radio(
        "Action",
        ["Translate (keep content identical)", "Rewrite natively (improve wording, keep facts)"],
    )
    focus = st.text_input("Optional focus (e.g. target role)", placeholder="e.g. Senior Data Engineer")
    run = st.button("✨ Generate", type="primary", use_container_width=True)

if run and resume_text.strip():
    client = get_client()
    with st.spinner(f"Generating {target_language} version…"):
        if action.startswith("Translate"):
            result = translate_resume(client, resume_text, LANGUAGES[target_language])
        else:
            result = rewrite_in_language(client, resume_text, LANGUAGES[target_language], focus)
    st.session_state["ml_result"] = result
    st.session_state["ml_lang"] = target_language

if st.session_state.get("ml_result"):
    st.divider()
    st.subheader(f"3. Result — {st.session_state.get('ml_lang', '')}")
    result = st.session_state["ml_result"]
    st.text_area("Generated text", result, height=350)

    c1, c2, c3 = st.columns(3)
    with c1:
        st.download_button(
            "📄 Download .docx",
            to_docx(result),
            file_name=f"resume_{st.session_state.get('ml_lang', '').lower()}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    with c2:
        st.download_button(
            "📕 Download .pdf",
            to_pdf(result),
            file_name=f"resume_{st.session_state.get('ml_lang', '').lower()}.pdf",
            mime="application/pdf",
        )
    with c3:
        if st.button("🔄 Start another", use_container_width=True):
            st.session_state.pop("ml_result", None)
            st.rerun()

st.divider()
st.info("💡 Tip: run the **Analyze** page first to get your parsed resume text, then paste it here to create localized versions for international applications.")
